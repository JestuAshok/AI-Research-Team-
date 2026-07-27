import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from backend.config import REPORTS_DIR

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def generate_docx_report(session_id: str, topic: str, summary_data: dict, sources_data: dict) -> str:
    """
    Generates a professional DOCX research report and saves it to the reports folder.
    Returns the absolute path to the generated file.
    """
    file_name = f"research_report_{session_id}.docx"
    file_path = REPORTS_DIR / file_name
    
    doc = Document()
    
    # Page setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Typography Colors
    plum_color = RGBColor(74, 21, 75)       # #4A154B
    aubergine_color = RGBColor(94, 39, 80)  # #5E2750
    rose_color = RGBColor(196, 140, 179)    # #C48CB3
    charcoal_color = RGBColor(45, 45, 45)   # #2D2D2D
    
    # Configure styles
    styles = doc.styles
    
    # Title formatting
    title_style = styles['Title']
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(28)
    title_font.bold = True
    title_font.color.rgb = plum_color
    
    # Heading 1 formatting
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Arial'
    h1_font.size = Pt(18)
    h1_font.bold = True
    h1_font.color.rgb = aubergine_color
    
    # Heading 2 formatting
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Arial'
    h2_font.size = Pt(13)
    h2_font.bold = True
    h2_font.color.rgb = rose_color
    
    # Body Text formatting
    body_style = styles['Normal']
    body_font = body_style.font
    body_font.name = 'Arial'
    body_font.size = Pt(10.5)
    body_font.color.rgb = charcoal_color

    # ------------------ COVER PAGE ------------------
    # Add title and spacer
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(80)
    p_title.paragraph_format.space_after = Pt(10)
    run_title = p_title.add_run(f"AI Intelligence Report:\n{topic}")
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = plum_color
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(150)
    run_sub = p_sub.add_run("A Multi-Agent Collaboration and Fact-Verified Analysis")
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = charcoal_color
    
    # Metadata
    date_str = datetime.datetime.now().strftime("%B %d, %Y")
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.line_spacing = 1.3
    
    run_meta = p_meta.add_run("Author: ")
    run_meta.bold = True
    p_meta.add_run("AI Research Team (6-Agent Swarm)\n")
    
    run_meta2 = p_meta.add_run("Session ID: ")
    run_meta2.bold = True
    p_meta.add_run(f"{session_id}\n")
    
    run_meta3 = p_meta.add_run("Date: ")
    run_meta3.bold = True
    p_meta.add_run(f"{date_str}\n")
    
    run_meta4 = p_meta.add_run("Confidence Level: ")
    run_meta4.bold = True
    p_meta.add_run(f"{summary_data.get('confidence_score', 85.0)}%")
    
    doc.add_page_break()

    # ------------------ 1. EXECUTIVE SUMMARY ------------------
    doc.add_heading("1. Executive Summary", level=1)
    p_exec = doc.add_paragraph()
    p_exec.paragraph_format.space_after = Pt(12)
    run_exec = p_exec.add_run(summary_data.get("executive_summary", ""))
    
    # Callout table for Key Statistics
    stats = summary_data.get("key_statistics", [])
    if stats:
        doc.add_heading("Key Statistics & Metrics", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Shading Accent 1'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        hdr_cells[2].text = 'Implication'
        
        # Apply style to headers
        for cell in hdr_cells:
            set_cell_background(cell, "5E2750")  # Aubergine
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Populate rows
        for stat in stats:
            row_cells = table.add_row().cells
            row_cells[0].text = str(stat.get("metric", ""))
            row_cells[1].text = str(stat.get("value", ""))
            row_cells[2].text = str(stat.get("description", ""))
            
            # Subtly shade row cells
            for cell in row_cells:
                set_cell_background(cell, "F8F4EC")
                
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ------------------ 2. INTRODUCTION & SCOPE ------------------
    doc.add_heading("2. Introduction & Research Intent", level=1)
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(12)
    p_intro.add_run(
        f"This intelligence dossier examines the multidimensional topic: {topic}. "
        "The synthesis is compiled by a collaborative framework of six specialized AI agents running on LangGraph. "
        "The Coordinator agent analyzed the primary objective, mapped out sub-questions, and tasked the Researcher "
        "agent to query Tavily search indexes and the arXiv archive. The collected sources were filtered, deduplicated, "
        "and scrutinized by a Fact Verification Agent to filter bias, identify contradictions, and assign a confidence rating."
    )

    # ------------------ 3. RESEARCH FINDINGS ------------------
    doc.add_heading("3. Detailed Research Findings", level=1)
    findings = summary_data.get("findings", [])
    if findings:
        for index, finding in enumerate(findings, 1):
            doc.add_heading(f"3.{index} {finding.get('subtopic', '')}", level=2)
            p_find = doc.add_paragraph()
            p_find.paragraph_format.space_after = Pt(12)
            p_find.add_run(finding.get('details', ''))
    else:
        p_empty = doc.add_paragraph()
        p_empty.add_run("No detailed findings found.")
        
    doc.add_page_break()

    # ------------------ 4. ADVANTAGES & CHALLENGES ------------------
    doc.add_heading("4. Architectural Challenges & Strengths", level=1)
    
    doc.add_heading("Advantages of Agentic Research Synthesis", level=2)
    advs = summary_data.get("advantages", [])
    for a in advs:
        p_adv = doc.add_paragraph(style='List Bullet')
        run_bold = p_adv.add_run(f"{a.get('title', '')}: ")
        run_bold.bold = True
        p_adv.add_run(a.get('description', ''))
        
    doc.add_heading("Technical and Operational Challenges", level=2)
    challs = summary_data.get("challenges", [])
    for c in challs:
        p_chall = doc.add_paragraph(style='List Bullet')
        run_bold = p_chall.add_run(f"{c.get('title', '')}: ")
        run_bold.bold = True
        p_chall.add_run(c.get('description', ''))

    # ------------------ 5. CONCLUSION & FUTURE SCOPE ------------------
    doc.add_heading("5. Conclusion & Future Outlook", level=1)
    p_concl = doc.add_paragraph()
    p_concl.paragraph_format.space_after = Pt(12)
    p_concl.add_run(summary_data.get("conclusion", ""))
    
    doc.add_heading("Future Research Scope", level=2)
    p_future = doc.add_paragraph()
    p_future.paragraph_format.space_after = Pt(12)
    p_future.add_run(summary_data.get("future_scope", ""))

    # ------------------ 6. REFERENCES ------------------
    doc.add_heading("6. Verified References & Papers", level=1)
    
    papers = sources_data.get("papers", [])
    if papers:
        doc.add_heading("Academic Papers (arXiv)", level=2)
        for index, p in enumerate(papers, 1):
            p_ref = doc.add_paragraph()
            p_ref.paragraph_format.space_after = Pt(6)
            run_num = p_ref.add_run(f"[{index}] ")
            run_num.bold = True
            run_title = p_ref.add_run(f"{p.get('title')} ")
            run_title.bold = True
            p_ref.add_run(f"by {p.get('authors')}. Published: {p.get('published')}.\nURL: {p.get('url')}\n")
            run_abs = p_ref.add_run(f"Abstract: {p.get('summary')[:200]}...")
            run_abs.italic = True
            
    web_sources = sources_data.get("web_sources", [])
    if web_sources:
        doc.add_heading("Web References (Tavily)", level=2)
        for index, w in enumerate(web_sources, 1):
            w_ref = doc.add_paragraph()
            w_ref.paragraph_format.space_after = Pt(6)
            run_num = w_ref.add_run(f"[W{index}] ")
            run_num.bold = True
            run_title = w_ref.add_run(f"{w.get('title')} ")
            run_title.bold = True
            p_ref_link = w_ref.add_run(f"\nURL: {w.get('url')} | Credibility Score: {int(w.get('score', 0.8)*100)}%\n")
            run_content = w_ref.add_run(f"Snippet: {w.get('content')[:200]}...")
            run_content.italic = True

    doc.save(str(file_path))
    print(f"[SUCCESS] DOCX report successfully generated at: {file_path}")
    return str(file_path)
